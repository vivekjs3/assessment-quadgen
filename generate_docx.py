import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, '0F172A')
    set_cell_margins(cell, top=240, bottom=240, left=240, right=240)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('VIVEK POD PLATFORM PROJ')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Enterprise Kubernetes & GitOps CI/CD Platform — On-Call Runbook & Architecture Specification')
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()

    h1 = doc.add_heading('1. Executive Overview & Platform Specification', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    doc.add_paragraph(
        'The Vivek Pod Platform Proj is a production-grade, GitOps-driven CI/CD platform engineered to eliminate manual Jenkins configuration overhead. '
        'Built entirely around zero-UI infrastructure-as-code principles, all authentication, RBAC authorization matrices, plugin management, and cloud pod agent templates '
        'are version-controlled and seeded dynamically via Jenkins Configuration-as-Code (JCasC) and Job DSL.'
    )
    
    h1 = doc.add_heading('2. System Architecture & Component Design', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph('Below is the end-to-end component layout and data flow across the host environment and Kubernetes cluster:')

    arch_box = doc.add_table(rows=1, cols=1)
    cell_arch = arch_box.cell(0, 0)
    set_cell_background(cell_arch, 'F8FAFC')
    set_cell_margins(cell_arch, top=180, bottom=180, left=180, right=180)
    p_arch = cell_arch.paragraphs[0]
    p_arch.paragraph_format.space_before = Pt(0)
    p_arch.paragraph_format.space_after = Pt(0)
    
    diagram_text = (
        "+-----------------------------------------------------------------------------------+\n"
        "|                            VIVEK POD PLATFORM ARCHITECTURE                        |\n"
        "+-----------------------------------------------------------------------------------+\n"
        "|  [ Proxmox VM / Baremetal Host ]                                                  |\n"
        "|       |                                                                           |\n"
        "|       +--> KinD Kubernetes Cluster (\"kind-jenkins\")                               |\n"
        "|               |                                                                   |\n"
        "|               +--> Namespace: jenkins                                             |\n"
        "|               |       |--> Jenkins Controller Pod (Port 8080 / NodePort 30080)   |\n"
        "|               |       |--> Ephemeral Agent Pod (Dynamic per pipeline execution)  |\n"
        "|               |                 |--> jnlp container (Jenkins Agent)               |\n"
        "|               |                 |--> builder container (kubectl & tools)         |\n"
        "|               |                                                                   |\n"
        "|               +--> Namespace: sample-app                                          |\n"
        "|                       |--> Vivek Pod Platform Proj (Nginx App / NodePort 30081)  |\n"
        "|                                                                                   |\n"
        "|       +--> Local Container Registry (\"kind-registry\" on Port 5001)                |\n"
        "+-----------------------------------------------------------------------------------+"
    )
    run_arch = p_arch.add_run(diagram_text)
    run_arch.font.name = 'Courier New'
    run_arch.font.size = Pt(9.5)
    run_arch.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph()

    h2 = doc.add_heading('2.1 Design Boundary Rationale: JCasC vs. K8s Manifests', level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    doc.add_paragraph(
        'A critical architectural decision was separating Jenkins infrastructure configuration (JCasC) from application delivery manifests (Kubernetes/Helm):\n\n'
        '• JCasC (jcasc/jenkins.yaml): Manages Jenkins controller state, RBAC roles, installed plugins, and Kubernetes cloud agent pod definitions. This guarantees build environment repeatability.\n'
        '• K8s Manifests (k8s/ & helm/): Define application deployment specifications, replica counts, ConfigMaps, and NodePort service exposures. Developers can modify application code and runtime specs independently without modifying Jenkins system configuration.\n'
        '• Security Scoping: The Jenkins agent operates under a dedicated ServiceAccount (jenkins-agent) strictly scoped via RoleBinding to the sample-app namespace, preventing unprivileged cluster-wide modification.'
    )

    h1 = doc.add_heading('3. On-Call Engineer Operational Runbook', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph('This section provides step-by-step triage procedures for engineers on call, covering the 4 most common failure scenarios:')

    triage_table = doc.add_table(rows=5, cols=3)
    triage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Scenario & Symptom', 'Root Cause', 'Resolution Command / Procedure']
    for idx, h_text in enumerate(headers):
        c = triage_table.cell(0, idx)
        set_cell_background(c, '1E293B')
        set_cell_margins(c, top=140, bottom=140, left=140, right=140)
        p = c.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    rows_data = [
        ('A: Ephemeral Agent Pod Failing to Spawn', 'JNLP tunnel port 50000 misconfigured or memory quota exceeded on KinD node.', 'kubectl apply -f k8s/jenkins-service.yaml\nkubectl rollout restart deployment/jenkins -n jenkins'),
        ('B: Controller Init Container Error / CrashLoop', 'Plugin dependency conflict or JCasC YAML syntax parsing failure.', 'kubectl logs -n jenkins -l app.kubernetes.io/name=jenkins -c copy-plugins\nValidate jcasc/jenkins.yaml'),
        ('C: Pipeline RBAC 403 Forbidden', 'jenkins-agent ServiceAccount missing RoleBinding in target namespace.', 'kubectl apply -f k8s/jenkins-rbac.yaml\nVerify: kubectl auth can-i create deployments --as=system:serviceaccount:jenkins:jenkins -n sample-app'),
        ('D: ImagePullBackOff from Local Registry', 'kind-registry container disconnected from KinD Docker bridge network.', 'docker network connect kind kind-registry || true\nkubectl apply -f k8s/sample-app/deployment.yaml')
    ]

    for row_idx, data in enumerate(rows_data, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(data):
            c = triage_table.cell(row_idx, col_idx)
            set_cell_background(c, bg_color)
            set_cell_margins(c, top=100, bottom=100, left=100, right=100)
            p = c.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)

    doc.add_paragraph()

    h2 = doc.add_heading('3.1 Emergency Recovery (1-Command Bootstrap)', level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    doc.add_paragraph(
        'In the event of catastrophic cluster corruption, run the automated teardown and bootstrap scripts:\n\n'
        '1. Tear down broken cluster: make down  (or ./scripts/teardown.sh)\n'
        '2. Restore zero-state environment: make up  (or ./scripts/bootstrap.sh)\n\n'
        'All cluster nodes, container registries, RBAC policies, and pipelines will be fully provisioned automatically.'
    )

    doc.save('/root/Vivek_Pod_Platform_Architecture_and_Runbook.docx')
    print('SUCCESSFULLY GENERATED VIVEK POD PLATFORM DOCX!')

if __name__ == '__main__':
    create_document()
